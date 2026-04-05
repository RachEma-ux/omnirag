"""DOCX extractor — headings, paragraphs, tables, comments."""

from __future__ import annotations

import io

from omnirag.intake.extractors.base import BaseExtractor
from omnirag.intake.models import ExtractedContent, RawContent, SourceObject


class DocxExtractor(BaseExtractor):
    name = "docx"

    def can_handle(self, mime_type: str | None, extension: str | None) -> bool:
        ext = (extension or "").lower().lstrip(".")
        if ext in ("docx", "doc"):
            return True
        if mime_type and ("wordprocessingml" in mime_type or "msword" in mime_type):
            return True
        return False

    async def extract(self, raw: RawContent, source_object: SourceObject) -> list[ExtractedContent]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = Document(io.BytesIO(raw.data))
        paragraphs = []
        headings = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                headings.append({"level": para.style.name, "text": text})
            paragraphs.append(text)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append(rows)

        full_text = "\n\n".join(paragraphs)
        if tables:
            for i, t in enumerate(tables):
                table_text = "\n".join([" | ".join(row) for row in t])
                full_text += f"\n\n[Table {i + 1}]\n{table_text}"

        if not full_text.strip():
            return []

        return [ExtractedContent(
            text=full_text,
            structure={
                "type": "document",
                "headings": headings,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
            },
            metadata={"filename": raw.filename},
            confidence=0.95,
        )]
